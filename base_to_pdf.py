import sqlite3
from fpdf import FPDF
import unicodedata
from docx import Document
from docx.shared import Inches


def unicode_normalize(s):
    return unicodedata.normalize('NFKD', s).encode('ascii', 'ignore')


def base_data(file_name, table_name):
    conn = sqlite3.connect(file_name)
    cursor = conn.cursor()
    sql = 'select count(*) from ' + table_name
    cursor.execute(sql)
    sql = 'select count(*) from pragma_table_info("' + table_name + '")'
    n = cursor.fetchone()
    cursor.execute(sql)
    m = cursor.fetchone()
    sql = 'SELECT * FROM ' + table_name
    cursor.execute(sql)
    a = cursor.fetchall()
    return a, n, m


def table_to_pdf(file_name, table_name, s, ans):
    ans = bool(ans)
    matric = base_data(file_name, table_name)
    matric = list(map(list, matric))
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=s, ln=1, align="C")
    #.encode('UTF-5', 'replace')
    if ans:
        pdf.cell(0, 10, txt="Equality holds", ln=2, align="C")
    else:
        pdf.cell(0, 10, txt="No equality holds", ln=2, align="C")
    #print(''.join(file_name.split('.')[:-1]) + ".pdf")
    pdf.output(''.join(file_name.split('.')[:-1]) + ".pdf")


def table_to_docx(file_name, table_name, s, ans):
    ans = bool(ans)
    matric = base_data(file_name, table_name)
    matric, n, m = list(map(list, matric))
    document = Document()

    document.add_heading(s, 0)
    if ans:
        document.add_paragraph('Равенство выполняется')
    else:
        document.add_paragraph('Равенство не выполняется')
    table1 = document.add_table(rows=0, cols=int(m[0]))
    print(n, m)
    #hdr_cells=table1.rows[0]
    for i in matric:
        row_cells = table1.add_row().cells
        for j in range(len(i)):
            row_cells[j].text = i[j]
    table1.style = 'Table Grid'
    #document.save("testing.docx")
    document.save(''.join(file_name.split('.')[:-1]) + ".docx")






#table_to_docx('5.db', 'name_table1', 'A=A', 'True')
